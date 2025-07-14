import streamlit as st
import streamlit.components.v1 as components
from sklearn.feature_extraction.text import CountVectorizer
import numpy as np
import PyPDF2
import docx
import os
import openai
from dotenv import load_dotenv
import re
import openai
import streamlit as st
from pynput import mouse, keyboard
import streamlit.components.v1 as components
from fpdf import FPDF
import asyncio
import time
from pathlib import Path
import plotly.graph_objects as go
import uuid
from streamlit_msal import Msal
from time import sleep
import threading
from streamlit_js_eval import streamlit_js_eval
# import textract
from tempfile import NamedTemporaryFile
from spire.doc import Document
from streamlit_autorefresh import st_autorefresh
import streamlit.components.v1 as components
import asyncio
from pynput import mouse, keyboard
from threading import Thread
# from session_config import SessionConfig
# from session_manager import init_session, track_activity, check_timeouts
load_dotenv()

st.set_page_config(page_title="Resume Filter App", page_icon="nathcorp.jpg", layout="centered")
# st.sidebar.page_link("pages/home.py", label="Home")
# st.sidebar.page_link("pages/instruction.py", label="Instruction")

TIMEOUT_MINUTES = 1 # one minute for testing :) 
openai_api_key = os.getenv('OPENAI_API_KEY')
openai_model_name = os.getenv('OPENAI_MODEL_NAME', 'gpt-3.5-turbo')
# # # print("asdasd", openai_model_name)
is_weight_valid = True

if "access_token" not in st.session_state or not st.session_state.get("logged_in"):
    st.warning("You are not logged in. Redirecting to login page...")
    st.switch_page("auth.py")  # Change to correct entry point if needed
 


# config = SessionConfig(
#     absolute_timeout=60,
#     idle_timeout=30,
#     warning_before=10,
#     login_page="auth.py"
# )

# init_session()          # sets up timers & heartbeat
# track_activity()        # captures any mouse/keyboard/scroll/click
# check_timeouts(config)
  

# OpenAI API key
openai.api_key = openai_api_key
# # # print("openai key", openai_api_key)
 
# Helper function to extract text from PDF
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ''
    for page in reader.pages:
        text += page.extract_text()
    return text



INACTIVITY_TIMEOUT = 5 * 60  # 5 minutes
last_activity_time = time.time()
logged_out = False
    
    
 

def on_key_press(key):
    try:
        print(f"Key pressed: {key.char}")
    except AttributeError:
        print(f"Special key pressed: {key}")

def on_key_release(key):
    update_activity_timestamp()
    print(f"Key released: {key}")
    if key == keyboard.Key.esc:
        return False  # Stop listener
  


def update_activity_timestamp():
    global last_activity_time, logged_out
    last_activity_time = time.time()  # reset to current time
    logged_out = False  # reset logout flag



def on_mouse_move(x, y):
    update_activity_timestamp()
    print(f"Mouse moved to ({x}, {y})")
    
    
async def inactivity_monitor():
    global logged_out
    while True:
        elapsed = time.time() - last_activity_time
        print(f"⏱️ Elapsed inactivity time: {int(elapsed)} seconds", end="\r")

        if elapsed > INACTIVITY_TIMEOUT and not logged_out:
            st.session_state.clear()
            Msal.sign_out()
            st.switch_page("auth.py")
            logged_out = True

        await asyncio.sleep(1)
        
def start_listeners():
    mouse_listener = mouse.Listener(
        on_move=on_mouse_move,
    )
    keyboard_listener = keyboard.Listener(
        on_press=on_key_press,
        on_release=on_key_release
    )
    mouse_listener.start()
    keyboard_listener.start()
    mouse_listener.join()
    keyboard_listener.join()

async def monitor_activity_async():
    print("👀 Monitoring for activity... Press ESC to exit.")

    listener_thread = Thread(target=start_listeners, daemon=True)
    listener_thread.start()

    await inactivity_monitor()  # Runs alongside listener thread

# Helper function to extract text from Word document
def extract_text_from_docx(file):
    doc = docx.Document(file)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def extract_text_from_doc(uploaded_file) -> str:
    # Write upload to a temp .doc file
    with NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    # Load and extract text
    doc = Document()
    doc.LoadFromFile(tmp_path)
    return doc.GetText()

def validate_jd(jd_text: str) -> str:
     
    jd_text = jd_text.strip()

    # Check if empty
    if not jd_text:
        st.warning("The job description is empty. Please enter a valid job description.")
        return False

    # Check minimum word count
    if len(jd_text.split()) < 10:
        st.warning("The job description is too short. Please provide more details.")
        return False

    # Check length limit (optional)
    if len(jd_text) > 5000:
        st.warning("The job description is too long. Please limit it to 5000 characters.")
        return False

    return True  



def logout():
    st.session_state.clear()
    Msal.sign_out()
    st.switch_page("auth.py")
    


def extract_skills_from_jd(job_description):

    prompt = (
        f"Extract only skills required from this JD and dont add eny extra word apart from the JD text. Just give me a list of technology without any special characters."
        f"{job_description}"
    )

    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Parse the response content
        response_text = response['choices'][0]['message']['content']
        # # # print("text...  " + response_text)
        # Extract potential skills using regex and cleanup
        skills = re.split(r',|\n|-', response_text)  # Split by commas, newlines, or hyphens
        skills = [skill.strip() for skill in skills if skill.strip()]  # Remove extra whitespace
        return skills  # Limit to max_skills

    except Exception as e:
        # # print(f"Error extracting skills: {e}")
        return []


def extract_skills_from_resume(resume, max_skills=20):

    prompt = (
        f"Extract only skills keywords from the resume"
        f"{resume}"
    )

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Parse the response content
        response_text = response['choices'][0]['message']['content']
        # # # print("text in resume...  " + response_text)
        # Extract potential skills using regex and cleanup
        skills = re.split(r',|\n|-', response_text)  # Split by commas, newlines, or hyphens
        skills = [skill.strip() for skill in skills if skill.strip()]  # Remove extra whitespace
        return skills[:max_skills]  # Limit to max_skills

    except Exception as e:
        # # print(f"Error extracting skills: {e}")
        return []




def extract_experience_from_jd(job_description):
    
    prompt = f"""    
    Job Description:
    {job_description}
    
    Only give how much experience required in number.
    """
    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[{"role": "system", "content": "You are an expert job description analyzer."},
                      {"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        keywords = response["choices"][0]["message"]["content"].strip()
        return set(map(str.strip, keywords.split(',')))
    except Exception as e:
        # # print(f"Error in extracting experience keywords: {e}")
        return set()
    
    
def extract_experience_from_resume(resume_eperience):
    
    prompt = f"""    
    Resume:
    {resume_eperience}
    
    You are an AI assistant that extracts total years of experience from resume text.

    Step 1: Check if the given input looks like a resume. If it does not resemble a resume (e.g., it looks like release notes, software documentation, bug reports, or other non-resume text), return only:
    0

    Step 2: If it is a resume, extract each job title, company name, and duration. Then calculate the total professional experience, accounting for overlapping dates, and return the total experience in **years and months**.

    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        keywords = response["choices"][0]["message"]["content"].strip()
        return set(map(str.strip, keywords.split(',')))
    except Exception as e:
        # # print(f"Error in extracting experience keywords: {e}")
        return set()
 
def extract_education_from_resume(resume_education, openai_model_name="gpt-4o"):
            prompt = f"""
            You are an AI assistant that extracts the candidate's highest education qualification from resume text.

            Step 1: Check if the input looks like a resume. If it doesn't, return only:
            "None"

            Step 2: If it is a resume, identify the **highest education qualification** such as Bachelor's, Master's, or PhD, along with the field of study (if available). Do not return a list. Just summarize the highest education level in a single line.

            Resume:
            \"\"\"
            {resume_education}
            \"\"\"
            """

            try:
                response = openai.ChatCompletion.create(
                    model=openai_model_name,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0,
                    top_p=1,
                    frequency_penalty=0,
                    presence_penalty=0.1,
                    max_tokens=300
                )
                result = response["choices"][0]["message"]["content"].strip()
                return result
            except Exception as e:
                # print(f"Error extracting education: {e}")
                return "None"
    
def extract_education_from_jd(job_description):
    
    prompt = f"""
    What educations required for this job just give me the degrees.
    
    Job Description:
    {job_description}
    
    """
    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        keywords = response["choices"][0]["message"]["content"].strip()
        return set(map(str.strip, keywords.split(',')))
    except Exception as e:
        # # print(f"Error in extracting education keywords: {e}")
        return set()


def extract_soft_skills_from_jd(job_description):
    """
    Extract soft skills-related keywords from the Job Description using OpenAI.
    """
    prompt = f"""
    Extract all soft skills-related keywords or phrases from the following Job Description. Focus on terms like "communication", "teamwork", "leadership", "problem-solving", "creativity", or similar:
    
    Job Description:
    {job_description}
    
    Provide the keywords or phrases as a comma-separated list.
    """
    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[{"role": "system", "content": "You are an expert job description analyzer."},
                      {"role": "user", "content": prompt}],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        keywords = response["choices"][0]["message"]["content"].strip()
        return set(map(str.strip, keywords.split(',')))
    except Exception as e:
        # # print(f"Error in extracting soft skills: {e}")
        return set()

def calculate_skill_match_percentage(skill_keywords, resume_keywords):
    """
    Calculate the skill match percentage between job description (JD) skills and resume skills.

    :param skill_keywords: Set of required skills in the job description (JD).
    :param resume_keywords: Set of skills extracted from the resume.
    :return: Skill match percentage.
    """
    # Normalize resume keywords by filtering only those present in skill_keywords
    matched_skills = skill_keywords.intersection(resume_keywords)
    # # print("matched skill", matched_skills)
    # Calculate percentage
    total_skills_in_jd = len(skill_keywords)
    matching_skills_count = len(matched_skills)

    if total_skills_in_jd == 0:
        return 0.0  # Avoid division by zero

    match_percentage = (matching_skills_count / total_skills_in_jd) * 100

    return round(match_percentage, 2), matched_skills  # Returning rounded percentage


def resumeValidation(resume_text):
    """
    Validate the resume text to ensure it contains essential sections.
    
    :param resume_text: The text content of the resume.
    :return: True if valid, False otherwise.
    """
    # Check if text is from a valid resume 

    # Normalize text
    resume_text = resume_text.lower()

    # Keywords commonly found in resumes
    essential_keywords = [
        "education", "experience", "skills", "projects",
        "summary", "objective", "certifications", "contact", "phone", "email", "linkedin", "github", "portfolio"
    ]

    # Count how many keywords are found
    found_keywords = sum(1 for keyword in essential_keywords if keyword in resume_text)

    return found_keywords >= 2  # At least 2 essential sections should be present


def calculate_exp_match_score(experience_keywords, experience_from_resume):
    
    # Convert sets to comma-separated lists
    exp_from_jd = ", ".join(experience_keywords)
    resume_exp_str = ", ".join(experience_from_resume)

    prompt = f"""
    

    - **Experience from JD**: {exp_from_jd}
    - **Resume experience**: {resume_exp_str}
    
    Calculate how much the experience is fulfilled if the JD or Resume does not contain any experience keywords then return 0 else if the
    resume experience is more than the JD give 100 else calculate and give the percentage.
    And dont assume anything just give me the percentage and dont take any example.
    
    Method:
    Convert the years into months from the resume and calculate the percentage.
    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                
                {"role": "user", "content": prompt}
            ],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Extract and convert the response to float
        match_percentage = response["choices"][0]["message"]["content"].strip()
        return match_percentage

    except Exception as e:
        # # print(f"Error calculating skill match percentage: {e}")
        return 0.0  # Default to 0% in case of an error


def calculate_edu_match_score(edu_keywords, edu_from_resume):
    
    # Convert sets to comma-separated lists
    edu_from_jd = ", ".join(edu_keywords)
    resume_edu_str = ", ".join(edu_from_resume)

    prompt = f"""
    

    - **Education from JD**: {edu_from_jd}
    - **Resume Education**: {resume_edu_str}
    Calculate how much the education is fulfilled if the Resume or JD does not contain any education keywprds then return 0 else if the
    JD education is there in the resume education then give 100 else give 0
    and dont give score for the cgpa. Just give me the percentage. either 0 or 100.
    """

    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[
                
                {"role": "user", "content": prompt}
            ],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Extract and convert the response to float
        match_percentage = response["choices"][0]["message"]["content"].strip()
        return match_percentage

    except Exception as e:
        # # print(f"Error calculating skill match percentage: {e}")
        return 0.0  # Default to 0% in case of an error
        
        
def calculate_soft_skill_match_score(softskilljd, resumesoftskill):
    
    # Convert sets to comma-separated lists
    var_softskilljd = ", ".join(softskilljd)
    var_resumesoftskill = ", ".join(resumesoftskill)

    prompt = f"""
    

    - **Soft Skill from JD**: {var_softskilljd}
    - **Resume Soft Skill**: {var_resumesoftskill}
    Calculate how much the soft skill is matching with the resume soft skills and just give me the percentage.
    """

    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[
                
                {"role": "user", "content": prompt}
            ],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Extract and convert the response to float
        match_percentage = response["choices"][0]["message"]["content"].strip()
        return match_percentage

    except Exception as e:
        # # print(f"Error calculating soft skill match percentage: {e}")
        return 0.0  # Default to 0% in case of an error

    
def extract_percentage(text):
    match = re.search(r"(\d+\.\d+)%", text)
    return float(match.group(1)) if match else 0.0

def extract_percent_ai_model(text):
    

    prompt = f"""
    
    extract the final percentage from the text and return the final percentage only.
    And just say "The final percent is xx.yy%"
    The text: {text}
    """

    try:
        response = openai.ChatCompletion.create(
            model=openai_model_name,
            messages=[
                
                {"role": "user", "content": prompt}
            ],
            temperature = 0,  # Low randomness for deterministic results
            top_p = 1,  # Moderate nucleus sampling for controlled diversity
            frequency_penalty = 0,  # Slight penalty to avoid repetitive answers
            presence_penalty = 0.1,  # Ensures all skills are considered
            max_tokens = 500,  # Restrict output length for predictability
        )
        
        # Extract and convert the response to float
        match_percentage = response["choices"][0]["message"]["content"].strip()
        return match_percentage

    except Exception as e:
        # # print(f"Error finding percentage: {e}")
        return 0.0  # Default to 0% in case of an error
 

def calculate_ats_score(resume_text, job_description, criteria_weights):
    if not resume_text:
        return 0.0, "No common skills has matched", "The resume is blank or has unrelated data"
    skill_keywords = set(extract_skills_from_jd(job_description))
    # # print("skill from jd", skill_keywords)
    resume_keywords = set(extract_skills_from_resume(resume_text))
    # # print("skill from resumee", resume_keywords)

    # Normalize resume skills (clean extra text formatting)
    # normalized_resume_keywords = {skill.split(":")[-1].strip() for skill in resume_keywords}

    # Calculate match percentage
    skill_match_percentage, matchedSkills = calculate_skill_match_percentage(skill_keywords, resume_keywords)
    
    # matched_skill = calculate_skill_match_percentage(skill_keywords, normalized_resume_keywords)
    # # print(f"Skill Match Percentage: {skill_match_percentage}%")
    skill_final_score_prev = extract_percent_ai_model(skill_match_percentage)
    skill_final_score = extract_percentage(skill_final_score_prev)
    
    # # print("Skill.. ", skill_final_score)
    experience_keywords = extract_experience_from_jd(job_description)
    experience_from_resume = extract_experience_from_resume(resume_text)
    print("experience from resume...", experience_from_resume)
    experience_match_score = calculate_exp_match_score(experience_keywords, experience_from_resume)
    # # print("experience match", experience_match_score)
    exp_final_score_prev = extract_percent_ai_model(experience_match_score)
    # # print("expre final scor eprev..", exp_final_score_prev)
    exp_final_score = extract_percentage(exp_final_score_prev)
    # # print("exp...", exp_final_score)
    # # # print("experience..", exp_final_score_prev)
    education_keywords = extract_education_from_jd(job_description)
    education_from_resume = extract_education_from_resume(resume_text)
    edu_score = calculate_edu_match_score(education_keywords, education_from_resume)
    edu_final_score_prev = extract_percent_ai_model(edu_score)

    edu_final_score = extract_percentage(edu_final_score_prev)
    # # print("edu...", edu_final_score)
    soft_skills = extract_soft_skills_from_jd(job_description)
    soft_skill_score = calculate_soft_skill_match_score(soft_skills, resume_keywords)
    # # print("soft skill score.......", soft_skill_score)
    soft_skill_final_score = extract_percentage(soft_skill_score)
    # # # print("softskillssssss...", soft_skill_final_score)
    # soft_skills_match = soft_skills & resume_keywords
    # soft_skill_score = soft_skills_match
    # soft_skill_score = (len(soft_skills_match) / len(soft_skills)) * criteria_weights.get("Soft Skills", 0) if soft_skills else 0
    # # print("soft...", soft_skill_final_score)
    # common_skills = skill_keywords & resume_keywords
    # Sum the weighted scores


    skill_weight = criteria_weights.get("Skills", 0)
    exp_weight = criteria_weights.get("Experience", 0)
    edu_weight = criteria_weights.get("Education", 0)
    soft_skill_weight = criteria_weights.get("Soft Skills", 0)

    skill_percentage_final = skill_final_score * (skill_weight / 100)
    exp_percentage_final = exp_final_score * (exp_weight / 100)
    edu_percentage_final = edu_final_score * (edu_weight / 100)
    soft_skill_percentage_final = soft_skill_final_score * (soft_skill_weight / 100)
    # # print("soft_skill_percentage_final", soft_skill_percentage_final)
    total_score = skill_percentage_final + exp_percentage_final + edu_percentage_final + soft_skill_percentage_final

    final_score = round(total_score, 2)
    return round(final_score, 2), matchedSkills, skill_keywords, round(skill_percentage_final, 2), round(exp_percentage_final, 2), round(edu_percentage_final, 2), round(soft_skill_percentage_final, 2)

# Add this function to check API key validity
def validate_api_key(api_key):
    try:
        # Test the API key by making a simple request
        openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": "Test"}]
        )
        return True
    except openai.error.AuthenticationError:
        return False
    except Exception as e:
        # # print(f"Unexpected error: {e}")
        return False

def extract_text_from_txt(file_path):
    """
    Reads and returns the content of a .txt file.

    :param file_path: Path to the .txt file
    :return: Text content as a string
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        return content
    except FileNotFoundError:
        return "File not found."
    except Exception as e:
        return f"An error occurred: {e}"

                    
def generate_resume_insights(
    resume_text,
    selected_criteria,
    criteria_weights,
    skill_percentage_final,
    exp_percentage_final,
    edu_percentage_final,
    soft_skill_percentage_final
):
    # Build criteria description safely
    criteria_description = ', '.join(
        [f"{criterion} ({criteria_weights.get(criterion, 0)}%)" for criterion in selected_criteria]
    )
    
    prefix = (
        "BEGIN RESPONSE WITH:\n"
        "Resume Analysis:\n"
        "Do NOT use any asterisk (*) characters in your output.\n"
        "Use bold text for headings, but do not use asterisks.\n"
        "If you need bold, use <strong>...</strong> or plain uppercase.\n\n"
    )
    
    
    # Use .get() for each weight inside the prompt to avoid KeyError
    body = f"""Analyze the following resume content based on the selected criteria and their weightage: {criteria_description}.
    Resume content:
    {resume_text}

    And don't give the final score. Also calculate the soft skill percentage. These are the final scores obtained for the given weightage; show them as well in the headline.

    1. Skills: {skill_percentage_final} / {criteria_weights.get("Skills", 0)}
    2. Experience: {exp_percentage_final} / {criteria_weights.get("Experience", 0)}
    3. Education: {edu_percentage_final} / {criteria_weights.get("Education", 0)}
    4. Soft Skills: {soft_skill_percentage_final} / {criteria_weights.get("Soft Skills", 0)}

    
    Show the skills, experience, education, and soft skills, summary heading in bold. And do not use the asterisk mark anywhere in the response.
    
    
    This is the response format:

     
    Skills: {skill_percentage_final}/{criteria_weights.get("Skills", 0)} (XX%)    
        Strengths: The candidate demonstrates a strong skill set in programming languages (Java, SQL, JavaScript, HTML, C++), modern technologies (AWS, Git/Github, Figma), and frameworks (ReactJs, NextJs, Springboot, NodeJs). The inclusion of both frontend and backend technologies indicates a well-rounded skill set suitable for software development roles.
        
    Experience: {exp_percentage_final}/{criteria_weights.get("Experience", 0)} (XX%)
        Strengths: The candidate has completed a 2-month internship as a Web Developer, which provides some practical experience in web development and collaboration.

    Education: {edu_percentage_final}/{criteria_weights.get("Education", 0)} (XX%)
        Strengths: The candidate is currently pursuing a B.Tech in Computer Science Engineering with a commendable CGPA of 8.44. The educational background includes a solid foundation from high school, with good scores in both matriculation and higher secondary education.

    Soft Skills: {soft_skill_percentage_final}/{criteria_weights.get("Soft Skills", 0)} (XX%)
        Strengths: The candidate mentions strong collaboration and problem-solving skills developed during the internship.

    Summary:

        Skills: XX.XX
        Experience: XX.XX 
        Education: XX.XX 
        Soft Skills: XX.XX 

    Overall, the candidate has a strong foundation in skills and education but needs to enhance their experience and soft skills sections to present a more balanced profile for potential employers.
    """

    prompt = prefix + body
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        temperature=0,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0.1,
        max_tokens=500,
        messages=[
            {"role": "system", "content": "You are an expert job matching assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response['choices'][0]['message']['content']


# Streamlit UI
MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024  # 5 MB in bytes

st.markdown(
    """
    <div style="background-color: white; padding: 1rem; text-align: center;">
        <img src="https://nathcorp.com/wp-content/uploads/2020/01/NathcorpLogo-Text-side_400x53.png" width="300">
    </div>
    """,
    unsafe_allow_html=True
) 


st.title("Resume Filter App", anchor=False)
st.markdown("""
    <style>
        [data-testid="stSidebar"] {
            height: 96vh !important;
            width: 20rem;
            overflow-y: auto;
        }
    </style>
""", unsafe_allow_html=True)


user = st.session_state.get("name", "User")
st.sidebar.write("👋 Welcome,", user)
st.sidebar.title("Select Criteria and Enter Weightage") 

# Smaller text + red asterisk
st.sidebar.markdown(
    "<p style='font-size:0.8em; margin-top:-0.5em;'>Fields marked with <span style='color:white'>*</span> are mandatory.</p>",
    unsafe_allow_html=True
)

st.markdown("""
    <script>
    const footer = document.querySelector("footer");
    const prefersDark = window.matchMedia("(prefers-color-scheme: dark)").matches;

    if (footer) {
        if (prefersDark) {
            footer.style.color = "white";
            footer.style.backgroundColor = "#111";
        } else {
            footer.style.color = "black";
            footer.style.backgroundColor = "white";
        }
    }
    </script>
""", unsafe_allow_html=True)



criteria_options = ["Skills", "Experience", "Education", "Soft Skills"]
criteria_value = [30, 15, 40, 15]

# Tooltip messages for each criterion
criteria_tooltips = {
    "Skills": "Set how much Skills affect the final ATS score",
    "Experience": "Set how much Experience affect the final ATS score",
    "Education": "Set how much Education affect the final ATS score",
    "Soft Skills": "Set how much Soft Skills affect the final ATS score"
}

# selected_criteria = []
# criteria_weights = {}

# for i, criterion in enumerate(criteria_options):
#     weight = st.sidebar.number_input(
#         f"Weightage for {criterion} (%)",
#         value=int(criteria_value[i]),
#         min_value=0,
#         step=1,
#         format="%d",
#         help=criteria_tooltips.get(criterion, "")  # Add tooltip here
#     )
#     criteria_weights[criterion] = weight
#     selected_criteria.append(criterion)

selected_criteria = []
criteria_weights = {}


pattern = re.compile(r"^(100|[1-9]?[0-9])$")  # Matches 0–100 integers

for i, criterion in enumerate(criteria_options):
    input_str = st.sidebar.text_input(
        f"Weightage for {criterion} (%) *",
        value=str(int(criteria_value[i])),
        help=criteria_tooltips.get(criterion, "")
    )

    if pattern.match(input_str):
        weight = int(input_str)
        criteria_weights[criterion] = weight
        selected_criteria.append(criterion)
    else:
        st.sidebar.warning(f"Enter a valid integer between 0 and 100 for {criterion}")
        is_weight_valid = False


 

st.markdown("""
    <style>
    /* Base footer styling */
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #26273;  /* keep your existing bg */
        text-align: center;
        padding: 10px;
        font-size: 14px;
        border-top: 1px solid #ccc;
    }
    /* Light theme: make text dark */
    html[data-theme='light'] .footer {
        color: #000 !important;
    }
    /* Dark theme: make text light */
    html[data-theme='dark'] .footer {
        color: #fff !important;
    }
    </style>
    <div class="footer">
        © 2025 Nathcorp | All rights reserved | v1.3
    </div>
""", unsafe_allow_html=True)




total_weight = sum(criteria_weights.values())
if total_weight > 100:
    # st.sidebar.warning("The total weightage exceeds 100%! Please adjust the values.")
    is_weight_valid = False
elif total_weight < 100:
    st.sidebar.info("The total weightage is less than 100%. You can allocate more.")
    is_weight_valid = False
elif total_weight == 100:
    st.sidebar.success("The total weightage is 100%.")
    is_weight_valid = True
   

# Use this as the main layout
with st.sidebar:
    # st.write("👋 Welcome, User")

    # Create a vertical spacer using markdown + HTML
    st.markdown(
        """
        <div style='height:14vh;'></div>
        """,
        unsafe_allow_html=True
    )

    # Button at the bottom
    if st.button("Logout"):
        logout()

st.write(
    """
    <div style="padding: 0; margin: 0; font-size: 20px; font-weight: 600;">
      Enter Job Description: <strong style="color: red;">*</strong>
    </div>
    """,
    unsafe_allow_html=True
)


job_description = st.text_area(
    label="",  # Hide the default label since we're using a custom heading
    height=200,
    help="Enter the job description to analyze resumes against.",
    placeholder="Enter or paste the job description.",
    disabled=not is_weight_valid
)

# Ensure uploader_key is in session_state for file_uploader reset
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = str(uuid.uuid4())

# Add a reset resumes button above the file uploader
col1, col2 = st.columns([1, 1])
with col1:
    reset_resumes = st.button("Reset All Resumes", key="reset_resumes_btn")

if reset_resumes:
    st.session_state["uploader_key"] = str(uuid.uuid4())
    st.rerun()

st.write(
    """
    <div style="padding: 0; margin: 0; font-size: 20px; font-weight: 600;">
      Upload Resumes: <strong style="color: red;">*</strong>
    </div>
    """,
    unsafe_allow_html=True
)


uploaded_resumes = st.file_uploader(
    label="",
    type=['pdf', 'docx', 'doc', 'txt'],
    accept_multiple_files=True,
    key=st.session_state["uploader_key"]
)
# check_session_timeout()

# uploaded_resumes = ""

# with open("ind.html", "r") as f:
#     html_uploader = f.read()

# components.html(html_uploader, height=100)

def valid_resume(resume_text):
    # Check if text is from a valid resume
    if not resume_text or len(resume_text.strip()) < 50:
        st.warning("The resume text is too short or empty. Please upload a valid resume.")
        return False

    # Normalize text
    resume_text = resume_text.lower()

    # Keywords commonly found in resumes
    essential_keywords = [
        "education", "experience", "skills", "projects",
        "summary", "objective", "certifications", "contact", "phone", "email"
    ]

    # Count how many keywords are found
    found_keywords = sum(1 for keyword in essential_keywords if keyword in resume_text)

    if found_keywords < 2:
        st.warning("The uploaded file doesn't appear to be a valid resume. Required section is missing. Please upload a valid resume.")
        return False 

    # If all checks pass
    return True

summary_data = []

if validate_api_key(openai_api_key) and uploaded_resumes and validate_jd(job_description):
    if total_weight == 100 and job_description:
        if st.button("Submit"):
            # Placeholders for live-updating summary table and details
            summary_placeholder = st.empty()
            # We'll render each resume's detail immediately after it's ready
            # No need for a separate container

            summary_data = []

            for idx, uploaded_resume in enumerate(uploaded_resumes):
                # — File size check —
                uploaded_resume.seek(0, 2)
                if uploaded_resume.tell() > MAX_FILE_SIZE_BYTES:
                    st.warning(f"{uploaded_resume.name} exceeds 5 MB; please upload a smaller file.")
                    st.stop()
                uploaded_resume.seek(0)

                # — Extract text —
                mime = uploaded_resume.type
                if mime == "application/pdf":
                    resume_text = extract_text_from_pdf(uploaded_resume)
                elif mime.endswith("wordprocessingml.document"):
                    resume_text = extract_text_from_docx(uploaded_resume)
                elif mime == "text/plain":
                    resume_text = extract_text_from_txt(uploaded_resume)
                elif mime == "application/msword":
                    resume_text = extract_text_from_doc(uploaded_resume)
                else:
                    resume_text = None

                if not (resume_text and valid_resume(resume_text) and resumeValidation(resume_text)):
                    continue

                with st.spinner(f"Analyzing {uploaded_resume.name}…"):
                    skills_from_jd = extract_skills_from_jd(job_description)
                    (
                        final_score,
                        common_skills,
                        skill_kw,
                        skill_pct,
                        exp_pct,
                        edu_pct,
                        soft_pct
                    ) = calculate_ats_score(resume_text, job_description, criteria_weights)

                    insights = generate_resume_insights(
                        resume_text,
                        selected_criteria,
                        criteria_weights,
                        skill_pct,
                        exp_pct,
                        edu_pct,
                        soft_pct
                    )

                # — Update summary_data and live-update the table —
                summary_data.append({
                    "Resume Name":     uploaded_resume.name,
                    "Skill %":         f"{skill_pct:.2f}",
                    "Experience %":    f"{exp_pct:.2f}",
                    "Education %":     f"{edu_pct:.2f}",
                    "Soft Skill %":    f"{soft_pct:.2f}",
                    "Final ATS Score": f"{final_score:.0f}",
                })

                # On first processed resume, render the “Resume Analysis” header
                if idx == 0:
                    st.write("Resume Analysis")

                # Live-update the summary table under the header
                summary_placeholder.markdown(
                    "<div style='padding:10px 0; margin-top:20px; "
                    "font-size:22px; font-weight:600;'>Resume Summary Table</div>",
                    unsafe_allow_html=True
                )
                summary_placeholder.dataframe(summary_data, width=700)

                # — Now render this resume’s detailed cards —

                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#d4edda;border:1px solid #c3e6cb;"
                    f"color:#155724;margin-bottom:15px;'>"
                    f"Analysis Completed for <strong>{uploaded_resume.name}</strong>!"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Skills card
                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#cce5ff;border:1px solid #b8daff;"
                    f"color:#004085;margin-bottom:15px;font-size:18px;line-height:1.6;'>"
                    f"<div style='font-size:24px;font-weight:bold;margin-bottom:10px;'>Skills</div>"
                    f"<strong>Matched Skills:</strong> {', '.join(common_skills)}<br>"
                    f"<strong>Extracted from JD:</strong> {', '.join(skills_from_jd)}<br>"
                    f"<strong>Skill %:</strong> {skill_pct:.2f}%"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Experience card
                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#cce5ff;border:1px solid #b8daff;"
                    f"color:#004085;margin-bottom:15px;font-size:18px;line-height:1.6;'>"
                    f"<div style='font-size:24px;font-weight:bold;margin-bottom:10px;'>Experience</div>"
                    f"<strong>Experience %:</strong> {exp_pct:.2f}%"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Education card
                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#cce5ff;border:1px solid #b8daff;"
                    f"color:#004085;margin-bottom:15px;font-size:18px;line-height:1.6;'>"
                    f"<div style='font-size:24px;font-weight:bold;margin-bottom:10px;'>Education</div>"
                    f"<strong>Education %:</strong> {edu_pct:.2f}%"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Soft Skills card
                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#cce5ff;border:1px solid #b8daff;"
                    f"color:#004085;margin-bottom:15px;font-size:18px;line-height:1.6;'>"
                    f"<div style='font-size:24px;font-weight:bold;margin-bottom:10px;'>Soft Skills</div>"
                    f"<strong>Soft Skills %:</strong> {soft_pct:.2f}%"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # GPT Insights card
                st.markdown(
                    f"<div style='padding:15px;border-radius:10px;"
                    f"background-color:#cce5ff;border:1px solid #b8daff;"
                    f"color:#004085;margin-bottom:15px;font-size:18px;line-height:1.6;'>"
                    f"<div style='font-size:24px;font-weight:bold;margin-bottom:10px;'>GPT Insights</div>"
                    f"{insights}"
                    f"</div>",
                    unsafe_allow_html=True
                )

                # Final ATS Score card
                st.markdown(
                    f"<div style='padding:20px;border-radius:12px;"
                    f"background-color:#000080;color:white;text-align:center;"
                    f"font-size:1.8em;margin-bottom:10px;'>"
                    f"Final ATS Score for {uploaded_resume.name}: {final_score:.0f}%"
                    f"</div>",
                    unsafe_allow_html=True
                )

            # — After loop: generate PDF of summary and offer download —
            if summary_data:
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", "B", 16)
                pdf.cell(200, 10, "Resume Summary Report", ln=True, align="C")
                pdf.set_font("Arial", "", 12)
                pdf.ln(10)
                headers = list(summary_data[0].keys())
                for row in summary_data:
                    for h in headers:
                        pdf.cell(0, 10, f"{h}: {row[h]}", ln=True)
                    pdf.ln(5)
                pdf_path = "resume_summary.pdf"
                pdf.output(pdf_path)
                if Path(pdf_path).exists():
                    with open(pdf_path, "rb") as f:
                        st.download_button(
                            label="📄 Download Resume Summary PDF",
                            data=f,
                            file_name="resume_summary.pdf",
                            mime="application/pdf"
                        )
                        
                        
if __name__ == "__main__":
    asyncio.run(monitor_activity_async())